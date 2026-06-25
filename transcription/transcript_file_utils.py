"""Transcript file extraction helpers for uploaded TXT, PDF, and DOCX files."""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document

from transcription.sarvam_client import TranscriptionResult, TranscriptionSegment


SUPPORTED_TRANSCRIPT_EXTENSIONS = {"txt", "pdf", "docx"}


class TranscriptFileError(Exception):
    """Raised when an uploaded transcript cannot be read or validated."""


@dataclass(frozen=True, slots=True)
class ExtractedTranscript:
    """Extracted transcript text and renderable transcript result."""

    text: str
    result: TranscriptionResult


def extract_uploaded_transcript(uploaded_file: BinaryIO, *, filename: str | None = None) -> ExtractedTranscript:
    """Extract transcript text from a TXT, PDF, or DOCX uploaded file."""

    display_name = filename or getattr(uploaded_file, "name", "") or "uploaded transcript"
    suffix = Path(display_name).suffix.lower().lstrip(".")
    if suffix not in SUPPORTED_TRANSCRIPT_EXTENSIONS:
        raise TranscriptFileError(
            "Unsupported transcript file. Please upload a TXT, PDF, or DOCX file."
        )

    file_bytes = _read_uploaded_bytes(uploaded_file)
    if not file_bytes:
        raise TranscriptFileError("The uploaded transcript file is empty.")

    try:
        if suffix == "txt":
            text = _extract_txt(file_bytes)
        elif suffix == "pdf":
            text = _extract_pdf(file_bytes)
        else:
            text = _extract_docx(file_bytes)
    except TranscriptFileError:
        raise
    except Exception as exc:
        raise TranscriptFileError(
            "We could not read this transcript file. Please try another TXT, PDF, or DOCX file."
        ) from exc

    text = _normalize_text(text)
    if not text:
        raise TranscriptFileError(
            "No transcript text was found in the uploaded file. Please upload a file with selectable text."
        )

    result = build_transcription_result_from_text(text, source_filename=display_name)
    return ExtractedTranscript(text=text, result=result)


def build_transcription_result_from_text(text: str, *, source_filename: str = "") -> TranscriptionResult:
    """Create a TranscriptionResult compatible with the existing UI/export flow."""

    normalized_text = _normalize_text(text)
    segments = _parse_speaker_segments(normalized_text)
    return TranscriptionResult(
        transcript=normalized_text,
        language_code=None,
        segments=segments,
        raw_payload={
            "source": "uploaded_transcript",
            "filename": source_filename,
        },
    )


def _read_uploaded_bytes(uploaded_file: BinaryIO) -> bytes:
    if hasattr(uploaded_file, "getvalue"):
        return uploaded_file.getvalue()

    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)
    data = uploaded_file.read()
    return data if isinstance(data, bytes) else bytes(data)


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise TranscriptFileError("This TXT file uses an unsupported text encoding.")


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise TranscriptFileError(
            "PDF transcript support is not installed. Please install project requirements and try again."
        ) from exc

    reader = PdfReader(BytesIO(file_bytes))
    page_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            page_text.append(text)
    return "\n\n".join(page_text)


def _extract_docx(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


_SPEAKER_LINE_RE = re.compile(
    r"^\s*(?P<label>(?:speaker|spk)\s*[\w.-]*)"
    r"(?:\s*\[(?P<start>[^\]-]+)\s*-\s*(?P<end>[^\]]+)\])?"
    r"\s*:?\s*(?P<rest>.*)$",
    re.IGNORECASE,
)


def _parse_speaker_segments(text: str) -> list[TranscriptionSegment]:
    segments: list[TranscriptionSegment] = []
    current_label: str | None = None
    current_start: float | None = None
    current_end: float | None = None
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_label, current_start, current_end, current_lines
        transcript = " ".join(line.strip() for line in current_lines if line.strip()).strip()
        if current_label and transcript:
            segments.append(
                TranscriptionSegment(
                    transcript=transcript,
                    speaker_id=_speaker_id_from_label(current_label),
                    start_time_seconds=current_start,
                    end_time_seconds=current_end,
                )
            )
        current_label = None
        current_start = None
        current_end = None
        current_lines = []

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        match = _SPEAKER_LINE_RE.match(stripped)
        if match:
            flush()
            current_label = match.group("label").strip()
            current_start = _parse_timestamp(match.group("start"))
            current_end = _parse_timestamp(match.group("end"))
            rest = match.group("rest").strip()
            if rest:
                current_lines.append(rest)
            continue

        if current_label:
            current_lines.append(stripped)

    flush()
    return segments


def _speaker_id_from_label(label: str) -> str:
    number_match = re.search(r"(\d+)", label)
    if number_match:
        number = int(number_match.group(1))
        return str(max(0, number - 1))
    return re.sub(r"\s+", " ", label).strip()


def _parse_timestamp(value: str | None) -> float | None:
    if not value:
        return None

    value = value.strip()
    parts = value.split(":")
    try:
        if len(parts) == 2:
            minutes, seconds = parts
            return int(minutes) * 60 + float(seconds)
        if len(parts) == 3:
            hours, minutes, seconds = parts
            return int(hours) * 3600 + int(minutes) * 60 + float(seconds)
        return float(value)
    except ValueError:
        return None
